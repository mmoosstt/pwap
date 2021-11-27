from spos.storage.db.setup import setupDatabase
from spos.config import config
from spos.storage.db.pupil import orm as ormP
from spos.storage.db.pupilvaluationset import orm as ormPVS
from spos.storage.db.pupilvaluation import orm as ormPV
from spos.storage.db.schoolclass import orm as ormSC
from pony.orm import db_session, desc

setupDatabase(config['database'])

@db_session
def run(document):
    for pupil in ormP.Pupil.select():
        sc = ormSC.SchoolClass.select(lambda sc:sc.id==pupil.schoolClass.id).first()
        document.add_heading(f'{pupil.givenName} {pupil.familyName}', 0)
        
        for pvs in ormPVS.PupilValuationSet.select(lambda pvs:pvs.pupil.id==pupil.id).order_by(desc(ormPVS.PupilValuationSet.changeDate)):
            date = pvs.changeDate
            document.add_heading(f'Bewertungen vom {pvs.changeDate.day}.{pvs.changeDate.month}.{pvs.changeDate.year}', 1)
            
            for pv in ormPV.PupilValuation.select(lambda pv: pv.pupilValuationSet.id == pvs.id).order_by(ormPV.PupilValuation.name):
                
                _translate = {
                    "attention":"Aufmerksamkeit",
                    "care":"Verhalten",
                    "relayiability":"Zuverlässigkeit",
                    "socialBehaviour":"Sozialverhalten",
                    "speaches":"Wortmeldungen",
                    "willingnessToLearn":"Lernbereitschaft",
                    }
                
                document.add_heading(f'{_translate[pv.name]}', 2)
                document.add_paragraph(f'{pv.gradeText}')


from docx import Document
from docx.shared import Inches

document = Document()

run(document)

document.save('demo.docx')