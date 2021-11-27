from spos.storage.db.setup import setupDatabase
from spos.config import config
from spos.storage.db.pupil import orm as ormP
from spos.storage.db.pupilvaluationset import orm as ormPVS
from spos.storage.db.pupilvaluation import orm as ormPV
from spos.storage.db.schoolclass import orm as ormSC
from pony.orm import db_session, desc
from docx import Document

import io

class Reports(object):
    
    def reportPupil(self, pupilId):
        document = Document()   
        document = self._reportPupil(pupilId, document, 0)
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream 

    def reportSchoolClass(self, schoolClassId):
        document = Document()   
        document = self._reportSchoolClass(schoolClassId, document, 0)
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream 
          
    @db_session
    def _reportPupil(self, pupilId, document, level=0):
        for pupil in ormP.Pupil.select(lambda p:p.id==pupilId):
            sc = ormSC.SchoolClass.select(lambda sc:sc.id==pupil.schoolClass.id).first()
            document.add_heading(f'{pupil.givenName} {pupil.familyName}', level)
            
            for pvs in ormPVS.PupilValuationSet.select(lambda pvs:pvs.pupil.id==pupil.id).order_by(desc(ormPVS.PupilValuationSet.changeDate)):
                date = pvs.changeDate
                document.add_heading(f'Bewertungen vom {pvs.changeDate.day}.{pvs.changeDate.month}.{pvs.changeDate.year}', level+1)
                
                for pv in ormPV.PupilValuation.select(lambda pv: pv.pupilValuationSet.id == pvs.id).order_by(ormPV.PupilValuation.name):
                    _translate = {
                        "attention":"Aufmerksamkeit",
                        "care":"Verhalten",
                        "relayiability":"Zuverlässigkeit",
                        "socialBehaviour":"Sozialverhalten",
                        "speaches":"Wortmeldungen",
                        "willingnessToLearn":"Lernbereitschaft",
                        "economics": 'Wirtschaft',
                        "english": 'Englisch',
                        }
                    document.add_heading(f'{_translate[pv.name]}', level+2)
                    document.add_paragraph(f'{pv.gradeText}')
                       
            return document
        
    @db_session
    def _reportSchoolClass(self, schoolClassId, document, level=0):
        
        for schoolClass in ormSC.SchoolClass.select(lambda sc:sc.id==schoolClassId):
            document.add_heading(f'Schulklasse - {schoolClass.name}', level)
            
            for pupil in schoolClass.pupils:
                document = self._reportPupil(pupil.id, document, level+1)
            
        return document



